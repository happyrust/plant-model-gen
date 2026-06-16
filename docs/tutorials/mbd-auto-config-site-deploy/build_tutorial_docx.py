import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SCREENSHOTS = ROOT / "screenshots"
OUTPUT = Path(
    os.environ.get(
        "MBD_TUTORIAL_OUTPUT",
        ROOT / "mbd-auto-config-site-deploy-tutorial.docx",
    )
)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color="2E74B5")
    elif level == 2:
        set_run_font(run, size=13, bold=True, color="2E74B5")
    else:
        set_run_font(run, size=12, bold=True, color="1F4E79")
    return paragraph


def add_body(doc, text="", bold_prefix=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        set_run_font(prefix, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Cm(0.6)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_code_line(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.5)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9.5)
    return paragraph


def add_screenshot(doc, filename, caption):
    image_path = SCREENSHOTS / filename
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.4))

    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run(caption)
    set_run_font(caption_run, size=9, color="666666")
    caption_p.paragraph_format.space_after = Pt(10)


def add_step(doc, number, title, body, screenshot, caption, checks=None):
    add_heading(doc, f"步骤 {number}: {title}", 2)
    for line in body:
        add_body(doc, line)
    if checks:
        for item in checks:
            add_bullet(doc, item)
    add_screenshot(doc, screenshot, caption)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("MBD 名称自动配置站点部署教程")
    set_run_font(title_run, size=20, bold=True, color="1F4E79")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("适用于 Plant Admin 的快速部署与站点编辑流程")
    set_run_font(subtitle_run, size=11, color="666666")

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    data = [
        ("目标 DB 文件", r"D:\AVEVA\Projects\E3D2.1\AvevaPlantSample\aps000\aps250160_0001"),
        ("MBD 名称", "ALL 或 /ALL"),
        ("搜索根目录", r"D:\AVEVA\Projects\E3D2.1"),
        ("教程生成日期", "2026-06-16"),
    ]
    for row, (key, value) in zip(table.rows, data):
        set_cell_shading(row.cells[0], "E8EEF5")
        row.cells[0].text = key
        row.cells[1].text = value
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=10, bold=(cell is row.cells[0]))

    add_heading(doc, "教程目标", 1)
    add_body(
        doc,
        "本教程演示如何只指定目标 DB 文件和 MBD 名称 ALL，让站点部署流程自动补齐关联工程路径，并在站点编辑里确认配置结果。",
    )
    add_bullet(doc, r"目标文件为 D:\AVEVA\Projects\E3D2.1\AvevaPlantSample\aps000\aps250160_0001。")
    add_bullet(doc, r"工程根目录为 D:\AVEVA\Projects\E3D2.1，系统会在该目录下发现 AvevaPlantSample、AvevaCatalogue 等关联工程。")
    add_bullet(doc, "保存或快速创建后，应能在解析范围里看到 dbnum=250160，并在工程组成里看到关联工程路径。")

    add_heading(doc, "前置条件", 1)
    add_bullet(doc, "Plant Admin 服务已启动，并且当前账号可以进入站点管理页面。")
    add_bullet(doc, "本机能够访问 E3D 工程目录，路径大小写和盘符与实际机器一致。")
    add_bullet(doc, "如果搜索根目录留空，后端会优先根据目标 DB 文件路径推断工程父目录；显式填写搜索根目录更利于排查。")

    add_step(
        doc,
        1,
        "登录 Plant Admin",
        [
            "打开管理入口，输入管理员账号与密码后登录。",
            "登录成功后会进入 Plant Admin 工作台，后续操作都在站点管理页面完成。",
        ],
        "01-login.png",
        "图 1 登录 Plant Admin。",
    )

    add_step(
        doc,
        2,
        "进入站点管理",
        [
            "点击左上方导航中的“站点管理”。页面上方会显示“快速创建部署”，下方是当前站点列表。",
            "这一页可以直接做 MBD 快速创建，也可以进入“新建站点/编辑站点”抽屉做更细的配置。",
        ],
        "02-sites-overview.png",
        "图 2 站点管理首页。",
    )

    add_step(
        doc,
        3,
        "使用快速创建部署的 MBD 模式",
        [
            "在“快速创建部署”区域选择“按 MBD 名称”。",
            "MBD 输入 ALL 或 /ALL；搜索根目录输入 D:\\AVEVA\\Projects\\E3D2.1。",
            "保持“自动解析依赖库”开启。需要按引用闭包解析 CATA 时，保持“按需解析 CATA”开启。",
        ],
        "03-quick-deploy-mbd-mode.png",
        "图 3 快速创建部署切换到按 MBD 名称。",
        checks=[
            "推荐输入 /ALL，可读性更接近 MDB 路径；后端会规范化为 ALL。",
            "搜索根目录应是多个 E3D 工程的父目录，而不是某个 dbfile 文件本身。",
        ],
    )

    add_step(
        doc,
        4,
        "打开站点编辑器",
        [
            "点击“新建站点”进入站点编辑抽屉。抽屉内包含项目信息、MBD 自动配置、工程组成、运行配置和解析范围。",
            "如果已有站点，也可以点击对应行的“查看/编辑配置”进入同一类表单。",
        ],
        "04-site-editor-basic.png",
        "图 4 站点编辑器包含 MBD 自动配置区域。",
    )

    add_step(
        doc,
        5,
        "填写目标 DB 文件与 MBD 名称",
        [
            "项目名称填写 AvevaPlantSample，项目路径填写目标 DB 文件完整路径。",
            "MBD 名称填写 ALL，搜索根目录填写 D:\\AVEVA\\Projects\\E3D2.1。",
            "项目代码按当前部署约定填写；示例中使用 1。",
        ],
        "05-site-editor-mbd-values.png",
        "图 5 在站点编辑器中填写目标 DB 与 MBD 信息。",
    )

    add_step(
        doc,
        6,
        "准备配置并检查表单回填",
        [
            "点击“MBD 自动配置”区域里的“准备配置”。",
            "表单会把 MBD 名称和目标 DB 合并到当前配置，并提示保存时会按 MDB 名称自动发现关联工程。",
            "如果项目名称已存在，页面会提示重名；这不影响验证自动配置能力，但正式保存前需要改成唯一名称。",
        ],
        "06-site-editor-prepared.png",
        "图 6 准备配置后，表单提示保存时会自动发现关联工程。",
    )

    add_step(
        doc,
        7,
        "在站点列表确认目标 DB 范围",
        [
            "回到站点列表，用搜索框筛选刚创建或刚更新的站点。",
            "确认状态摘要中出现“本次按当前范围解析：dbnum=250160”。这说明目标 DB 文件已正确反推出 dbnum。",
        ],
        "07-site-list-filtered-result.png",
        "图 7 站点列表显示 dbnum=250160。",
    )

    add_step(
        doc,
        8,
        "在编辑器确认关联工程",
        [
            "点击“查看/编辑配置”，在“工程组成”区域检查自动发现的工程路径。",
            "验证样例中可看到 AvevaCatalogue 与 AvevaPlantSample\\aps000，同时手动 DB Nums 保留 250160。",
        ],
        "08-site-editor-existing-config.png",
        "图 8 编辑器中可确认关联工程和解析范围。",
    )

    add_heading(doc, "验证结果", 1)
    add_body(doc, "本次验证覆盖了快速部署与站点编辑两条路径：")
    add_bullet(doc, "快速部署接口返回 202，并解析出 site_id=codex-mbd-all-fallback-221948-8089、dbnum=250160、resolved_db_file=aps250160_0001。")
    add_bullet(doc, "站点编辑创建接口返回 201，保存后再次更新接口返回 200。")
    add_bullet(doc, r"编辑器结果中保留了 AvevaCatalogue 与 D:\AVEVA\Projects\E3D2.1\AvevaPlantSample\aps000 两个关联工程，并将手动 DB Nums 设置为 250160。")

    add_heading(doc, "排查提示", 1)
    add_body(doc, "如果没有找到关联工程，优先检查下面几项：")
    add_bullet(doc, "搜索根目录是否指向 D:\\AVEVA\\Projects\\E3D2.1 这一层，而不是指到 AvevaPlantSample 或 aps000 目录内部。")
    add_bullet(doc, "目标 DB 文件名是否包含可识别的 dbnum，例如 aps250160_0001 对应 250160。")
    add_bullet(doc, "正式保存前项目名称必须唯一；重名只会阻止保存，不代表 MBD 自动配置失败。")
    add_bullet(doc, "如果 MBD ALL 中部分成员没有 NUMBDB，系统仍会优先使用目标 DB 文件作为部署目标，并继续发现同级工程依赖。")

    add_heading(doc, "推荐输入模板", 1)
    add_body(doc, "快速部署或站点编辑时可直接按以下值填写：")
    add_code_line(doc, r"目标 DB 文件: D:\AVEVA\Projects\E3D2.1\AvevaPlantSample\aps000\aps250160_0001")
    add_code_line(doc, "MBD 名称: ALL")
    add_code_line(doc, r"搜索根目录: D:\AVEVA\Projects\E3D2.1")
    add_code_line(doc, "手动 DB Nums: 250160")

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_doc())
